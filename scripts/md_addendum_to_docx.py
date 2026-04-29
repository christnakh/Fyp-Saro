#!/usr/bin/env python3
"""Convert SECOND_SEMESTER_ADDENDUM.md to SECOND_SEMESTER_ADDENDUM.docx (requires python-docx)."""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Install: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def add_inline_runs(paragraph, text: str) -> None:
    if not text:
        return
    chunks = re.split(r"(\*\*.+?\*\*|`[^`\n]+`)", text)
    for chunk in chunks:
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
            continue
        if chunk.startswith("`") and chunk.endswith("`"):
            r = paragraph.add_run(chunk[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(10)
            continue
        sub = re.split(r"(\*[^*\n]+\*)", chunk)
        for s in sub:
            if not s:
                continue
            if s.startswith("*") and s.endswith("*") and len(s) > 2:
                paragraph.add_run(s[1:-1]).italic = True
            else:
                paragraph.add_run(s)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def row_is_separator(cells: list[str]) -> bool:
    if not cells:
        return True
    for c in cells:
        t = c.replace(" ", "")
        if not t or not re.fullmatch(r"[\-:]+", t):
            return False
    return True


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    md_path = root / "SECOND_SEMESTER_ADDENDUM.md"
    out_path = root / "SECOND_SEMESTER_ADDENDUM.docx"

    if not md_path.exists():
        print(f"Missing: {md_path}", file=sys.stderr)
        sys.exit(1)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    i = 0
    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        r = p.add_run("\n".join(code_lines))
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        code_lines = []

    def flush_table(rows: list[list[str]]) -> None:
        if not rows:
            return
        data_rows = [r for r in rows if not row_is_separator(r)]
        if not data_rows:
            return
        ncols = max(len(r) for r in data_rows)
        table = doc.add_table(rows=len(data_rows), cols=ncols)
        table.style = "Table Grid"
        for ri, row in enumerate(data_rows):
            for ci in range(ncols):
                cell_text = row[ci] if ci < len(row) else ""
                cell_text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
                cell_text = re.sub(r"`([^`]+)`", r"\1", cell_text)
                table.rows[ri].cells[ci].text = cell_text
        doc.add_paragraph()

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if is_table_row(line):
            table_rows: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            flush_table(table_rows)
            continue

        stripped = line.strip()
        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("#"):
            hashes = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[hashes:].strip()
            level = min(max(hashes, 1), 3)
            doc.add_heading(text, level=level)
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        # Markdown bullet: "- item" (use hyphen+space in .md to avoid * / italic clashes)
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:].strip())
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    flush_code()
    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
